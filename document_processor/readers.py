"""
Document Readers

Handles reading and text extraction from various document formats.
"""

import logging
from pathlib import Path
from typing import Optional
import io

class DocumentReader:
    """Unified document reader for multiple file formats"""

    def __init__(self, config):
        """
        Initialize document reader

        Args:
            config: Application configuration
        """
        self.config = config
        self.logger = logging.getLogger(__name__)

    def read_document(self, file_path: Path) -> Optional[str]:
        """
        Read text content from a document file

        Args:
            file_path: Path to the document

        Returns:
            Extracted text content or None if failed
        """
        if not file_path.exists():
            self.logger.error(f"File does not exist: {file_path}")
            return None

        extension = file_path.suffix.lower()

        try:
            if extension == '.txt':
                return self._read_text_file(file_path)
            elif extension == '.pdf':
                return self._read_pdf_file(file_path)
            elif extension in ['.doc', '.docx']:
                return self._read_word_file(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._read_excel_file(file_path)
            else:
                self.logger.warning(f"Unsupported file type: {extension}")
                return None

        except Exception as e:
            self.logger.error(f"Failed to read {file_path}: {e}")
            return None

    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text file"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                self.logger.debug(f"Successfully read text file with {encoding} encoding")
                return content
            except UnicodeDecodeError:
                continue

        raise Exception(f"Could not decode text file with any supported encoding")

    def _read_pdf_file(self, file_path: Path) -> str:
        """Read PDF file"""
        try:
            import PyPDF2
        except ImportError:
            try:
                import pypdf as PyPDF2  # Alternative import name
            except ImportError:
                self.logger.error("PyPDF2 or pypdf not installed. Cannot read PDF files.")
                return ""

        try:
            text_content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from page {page_num} in {file_path}: {e}")

            content = '\n\n'.join(text_content)

            if not content.strip():
                self.logger.warning(f"No text content extracted from PDF: {file_path}")

            return content

        except Exception as e:
            self.logger.error(f"Failed to read PDF file {file_path}: {e}")
            return ""

    def _read_word_file(self, file_path: Path) -> str:
        """Read Word document (.doc/.docx)"""
        extension = file_path.suffix.lower()

        if extension == '.docx':
            return self._read_docx_file(file_path)
        elif extension == '.doc':
            return self._read_doc_file(file_path)
        else:
            raise ValueError(f"Unsupported Word file extension: {extension}")

    def _read_docx_file(self, file_path: Path) -> str:
        """Read .docx file"""
        try:
            from docx import Document
        except ImportError:
            self.logger.error("python-docx not installed. Cannot read .docx files.")
            return ""

        try:
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))

            content = '\n'.join(text_content)
            return content

        except Exception as e:
            self.logger.error(f"Failed to read .docx file {file_path}: {e}")
            return ""

    def _read_doc_file(self, file_path: Path) -> str:
        """Read legacy .doc file"""
        try:
            import textract
        except ImportError:
            self.logger.error("textract not installed. Cannot read .doc files.")
            return ""

        try:
            # Use textract to extract text from .doc files
            text = textract.process(str(file_path))
            return text.decode('utf-8')

        except Exception as e:
            self.logger.error(f"Failed to read .doc file {file_path}: {e}")
            return ""

    def _read_excel_file(self, file_path: Path) -> str:
        """Read Excel file (.xlsx/.xls)"""
        try:
            import pandas as pd
        except ImportError:
            self.logger.error("pandas not installed. Cannot read Excel files.")
            return ""

        try:
            # Read all sheets from the Excel file
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl' if file_path.suffix == '.xlsx' else None)

            text_content = []

            for sheet_name, df in excel_data.items():
                # Add sheet header
                text_content.append(f"=== {sheet_name} ===")

                # Convert DataFrame to text representation
                # Replace NaN values with empty strings
                df = df.fillna('')

                # Add column headers
                if not df.empty:
                    headers = ' | '.join(str(col) for col in df.columns)
                    text_content.append(headers)
                    text_content.append('-' * len(headers))

                    # Add rows
                    for _, row in df.iterrows():
                        row_text = ' | '.join(str(value) for value in row.values if str(value).strip())
                        if row_text.strip():
                            text_content.append(row_text)

                text_content.append('')  # Empty line between sheets

            content = '\n'.join(text_content)
            return content

        except Exception as e:
            self.logger.error(f"Failed to read Excel file {file_path}: {e}")
            return ""

    def validate_file(self, file_path: Path) -> bool:
        """
        Validate that a file can be processed

        Args:
            file_path: Path to the file

        Returns:
            bool: True if file can be processed
        """
        if not file_path.exists():
            return False

        if not file_path.is_file():
            return False

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.MAX_DOCUMENT_SIZE_MB:
            self.logger.warning(f"File too large ({file_size_mb:.1f}MB): {file_path}")
            return False

        # Check file extension
        extension = file_path.suffix.lower()
        if extension not in self.config.SUPPORTED_DOCUMENT_TYPES:
            self.logger.warning(f"Unsupported file type: {extension}")
            return False

        return True

    def get_file_info(self, file_path: Path) -> dict:
        """
        Get information about a file

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with file information
        """
        if not file_path.exists():
            return {"error": "File does not exist"}

        try:
            stat = file_path.stat()
            return {
                "name": file_path.name,
                "extension": file_path.suffix.lower(),
                "size_bytes": stat.st_size,
                "size_mb": stat.st_size / (1024 * 1024),
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "is_supported": file_path.suffix.lower() in self.config.SUPPORTED_DOCUMENT_TYPES,
                "can_process": self.validate_file(file_path)
            }
        except Exception as e:
            return {"error": str(e)}

class PDFReader:
    """Specialized PDF reader with advanced features"""

    @staticmethod
    def extract_text_with_layout(file_path: Path) -> str:
        """Extract text while preserving layout information"""
        try:
            import pdfplumber
        except ImportError:
            # Fallback to basic PDF reading
            reader = DocumentReader(None)
            return reader._read_pdf_file(file_path)

        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return '\n\n'.join(text_content)

        except Exception as e:
            logging.getLogger(__name__).error(f"Advanced PDF reading failed: {e}")
            # Fallback to basic method
            reader = DocumentReader(None)
            return reader._read_pdf_file(file_path)

    @staticmethod
    def extract_tables(file_path: Path) -> list:
        """Extract tables from PDF"""
        try:
            import pdfplumber
        except ImportError:
            return []

        try:
            tables = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)

            return tables

        except Exception as e:
            logging.getLogger(__name__).error(f"Table extraction failed: {e}")
            return []